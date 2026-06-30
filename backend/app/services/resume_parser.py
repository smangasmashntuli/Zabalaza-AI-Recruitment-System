import re
from datetime import datetime, timezone
from io import BytesIO
from typing import Dict, List, Optional, Tuple

import PyPDF2
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except Exception:
    pdfplumber = None
    PDFPLUMBER_AVAILABLE = False
import docx


class ResumeParser:
    """Extract structured information from CV files for profile auto-fill."""

    def __init__(self):
        self.email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        self.phone_pattern = r"(?:\+?\d[\d\s().-]{7,}\d)"
        self.location_pattern = r"^[A-Za-z\s]+(?:,\s*[A-Za-z\s]+)+$"
        self.date_range_pattern = re.compile(
            r"((?:19|20)\d{2})(?:\s*[-–]\s*((?:19|20)\d{2}|present|current))?",
            re.IGNORECASE,
        )
        self.url_pattern = re.compile(r"https?://[^\s]+", re.IGNORECASE)

        self.degree_patterns = [
            "bachelor", "master", "phd", "doctorate", "associate", "diploma", "advanced diploma",
            "b.s.", "postgraduate diploma", "m.s.", "b.a.", "m.a.", "bcom", "bsc", "msc", "mba",
        ]
        self.job_title_keywords = [
            "engineer", "developer", "analyst", "manager", "designer", "scientist",
            "consultant", "specialist", "administrator", "coordinator", "architect",
            "intern", "officer",
        ]
        self.section_headers = {
            "summary": [
                "summary", "professional summary", "profile summary", "overview", "about me",
                "professional profile", "career summary", "executive summary", "profile"
            ],
            "cover_letter": ["cover letter", "dear hiring manager", "dear recruiter"],
            "education": [
                "education", "academic background", "qualifications", "academic qualifications",
                "educational background", "academic history", "education & qualifications"
            ],
            "experience": [
                "experience", "work experience", "employment history", "professional experience",
                "work history", "professional background", "employment", "career history",
                "relevant experience", "professional experience"
            ],
            "skills": [
                "skills", "technical skills", "core competencies", "expertise",
                "technical competencies", "key skills", "professional skills",
                "competencies", "areas of expertise", "skill set"
            ],
            "certifications": [
                "certifications", "licenses", "courses", "certificates",
                "professional certifications", "credentials", "qualifications"
            ],
            "projects": [
                "projects", "portfolio", "publications", "open source",
                "personal projects", "key projects", "project experience",
                "notable projects", "side projects"
            ],
            "languages": [
                "languages", "language proficiency", "language skills",
                "spoken languages", "language abilities"
            ],
        }
        self.skill_keywords = [
            "python", "java", "javascript", "typescript", "react", "angular", "vue",
            "node", "express", "django", "flask", "fastapi", "sql", "nosql", "mongodb",
            "postgresql", "mysql", "docker", "kubernetes", "aws", "azure", "gcp",
            "machine learning", "deep learning", "nlp", "computer vision", "tensorflow",
            "pytorch", "scikit-learn", "pandas", "numpy", "data analysis", "statistics",
            "html", "css", "git", "agile", "scrum", "ci/cd", "devops", "rest api",
            "graphql", "microservices", "blockchain", "solidity", "rust", "go", "c++",
            "c#", ".net", "spring", "hibernate", "redis", "elasticsearch", "kafka",
        ]
        self.soft_skill_keywords = [
            "communication", "teamwork", "leadership", "problem solving", "critical thinking",
            "time management", "adaptability", "collaboration", "stakeholder management",
        ]
        self.language_keywords = [
            "english", "zulu", "xhosa", "afrikaans", "french", "spanish", "german", "arabic", "portuguese",
        ]

    def _normalize_date(self, value: str) -> str:
        """Normalize common date values to YYYY-MM (best-effort)."""
        if not value:
            return ""

        clean = value.strip().lower()
        if clean in {"present", "current", "now"}:
            return "present"

        year_match = re.search(r"((?:19|20)\d{2})", clean)
        year = year_match.group(1) if year_match else ""
        if not year:
            return ""

        month_map = {
            "jan": "01", "feb": "02", "mar": "03", "apr": "04", "may": "05", "jun": "06",
            "jul": "07", "aug": "08", "sep": "09", "oct": "10", "nov": "11", "dec": "12",
        }
        for token, month in month_map.items():
            if token in clean:
                return f"{year}-{month}"
        return f"{year}-01"

    def _normalize_phone(self, value: Optional[str]) -> Optional[str]:
        if not value:
            return None
        raw = value.strip()
        digits = re.sub(r"\D", "", raw)
        if not digits:
            return None

        if raw.startswith("+"):
            return f"+{digits}"
        if digits.startswith("00"):
            return f"+{digits[2:]}"
        if len(digits) == 10 and digits.startswith("0"):
            return f"+27{digits[1:]}"
        return f"+{digits}"

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        # Bound parsing to first 10 pages for speed and predictable processing.
        try:
            if PDFPLUMBER_AVAILABLE:
                with pdfplumber.open(BytesIO(file_content)) as pdf:
                    pages = pdf.pages[:10]
                    chunks = []
                    for page in pages:
                        page_text = page.extract_text() or ""
                        table_lines: List[str] = []
                        for table in page.extract_tables() or []:
                            for row in table:
                                row_text = " | ".join((cell or "").strip() for cell in row if cell is not None)
                                if row_text.strip():
                                    table_lines.append(row_text)
                        chunks.append("\n".join([page_text] + table_lines).strip())
                    return "\n".join(chunks)

            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            return "\n".join(page.extract_text() or "" for page in pdf_reader.pages[:10])
        except Exception as exc:
            print(f"Error extracting PDF text: {exc}")
            return ""

    def extract_text_from_docx(self, file_content: bytes) -> str:
        try:
            doc = docx.Document(BytesIO(file_content))
            lines: List[str] = []

            for paragraph in doc.paragraphs:
                if not paragraph.text:
                    continue
                style_name = (getattr(paragraph.style, "name", "") or "").lower()
                prefix = "- " if style_name.startswith("list") else ""
                lines.append(f"{prefix}{paragraph.text}".strip())

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))

            return "\n".join(lines)
        except Exception as exc:
            print(f"Error extracting DOCX text: {exc}")
            return ""

    def extract_text_from_txt(self, file_content: bytes) -> str:
        try:
            return file_content.decode("utf-8", errors="ignore")
        except Exception as exc:
            print(f"Error extracting TXT text: {exc}")
            return ""

    def extract_text(self, file_content: bytes, file_type: str) -> str:
        if file_type == "application/pdf":
            return self.extract_text_from_pdf(file_content)
        if file_type in [
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]:
            return self.extract_text_from_docx(file_content)
        if file_type in ["text/plain", "text/markdown"]:
            return self.extract_text_from_txt(file_content)
        return ""

    def _clean_lines(self, text: str) -> List[str]:
        return [line.strip() for line in text.splitlines() if line.strip()]

    def _header_matches(self, line: str, header: str) -> bool:
        header = header.strip().lower()
        if not header:
            return False
        normalized = line.lower().strip(":").strip()
        return bool(re.search(rf"\b{re.escape(header)}\b", normalized))

    def _get_section_lines(self, text: str, section_key: str) -> List[str]:
        lines = self._clean_lines(text)
        headers = self.section_headers.get(section_key, [])
        start_idx = None

        for idx, line in enumerate(lines):
            if any(self._header_matches(line, h) for h in headers):
                start_idx = idx + 1
                break

        if start_idx is None:
            return []

        collected = []
        for line in lines[start_idx:]:
            if any(any(self._header_matches(line, h) for h in values) for values in self.section_headers.values()):
                break
            collected.append(line)

        return collected

    def extract_name(self, text: str) -> Optional[str]:
        for line in self._clean_lines(text)[:6]:
            lowered = line.lower()
            if (
                len(line.split()) in (2, 3)
                and not re.search(self.email_pattern, line)
                and not re.search(r"\d", line)
                and "resume" not in lowered
                and "curriculum" not in lowered
            ):
                return line.title()
        return None

    def split_name(self, full_name: Optional[str]) -> Tuple[Optional[str], Optional[str]]:
        if not full_name:
            return None, None
        parts = full_name.split()
        if len(parts) == 1:
            return parts[0], None
        return parts[0], " ".join(parts[1:])

    def extract_email(self, text: str) -> Optional[str]:
        emails = re.findall(self.email_pattern, text)
        return emails[0] if emails else None

    def extract_phone(self, text: str) -> Optional[str]:
        candidates = re.findall(self.phone_pattern, text)
        for candidate in candidates:
            digits = re.sub(r"\D", "", candidate)
            if 9 <= len(digits) <= 15:
                return self._normalize_phone(candidate)
        return None

    def extract_location(self, text: str) -> Optional[str]:
        for line in self._clean_lines(text)[:10]:
            if re.search(self.location_pattern, line) and len(line) < 80 and "linkedin" not in line.lower():
                return line
        return None

    def extract_title(self, text: str) -> Optional[str]:
        for line in self._clean_lines(text)[1:8]:
            lowered = line.lower()
            if any(keyword in lowered for keyword in self.job_title_keywords) and len(line) < 80:
                return line
        return None

    def extract_skills(self, text: str) -> List[str]:
        text_lower = text.lower()
        text_norm = re.sub(r"[\W_]+", " ", text_lower)
        tokens = set(text_norm.split())

        found = set()
        for skill in self.skill_keywords:
            s = skill.lower()
            variants = {
                s,
                s.replace(" ", ""),
                s.replace(".", ""),
                s.replace(".", "").replace(" ", ""),
            }
            if any(v in text_lower or v in text_norm or v in tokens for v in variants):
                found.add(skill.title())

        return sorted(found)

    def extract_experience_years(self, text: str) -> float:
        patterns = [
            r"(\d+)\+?\s*years?\s+(?:of\s+)?experience",
            r"experience[:\s]+(\d+)\+?\s*years?",
        ]
        years = []
        for pattern in patterns:
            years.extend(int(match) for match in re.findall(pattern, text.lower()))

        if years:
            return float(max(years))

        date_ranges = self.date_range_pattern.findall(text)
        total = 0
        for start, end in date_ranges[:6]:
            end_year = int(end) if end and end.isdigit() else datetime.now(timezone.utc).year
            total += max(end_year - int(start), 0)
        return float(total) if total else 0.0

    def extract_education(self, text: str) -> List[Dict[str, str]]:
        education = []
        lines = self._get_section_lines(text, "education") or self._clean_lines(text)

        for line in lines:
            lowered = line.lower()
            matched_degree = next((degree for degree in self.degree_patterns if degree in lowered), None)
            if not matched_degree:
                continue

            year_matches = re.findall(r"(?:19|20)\d{2}", line)
            parts = [part.strip() for part in re.split(r"[-|,]", line) if part.strip()]
            school = "Not provided"
            if parts:
                year_like_parts = [idx for idx, part in enumerate(parts) if re.fullmatch(r"(?:19|20)\d{2}", part)]
                if year_like_parts:
                    candidate_idx = year_like_parts[0] - 1
                    if candidate_idx >= 0:
                        school = parts[candidate_idx]
                if school == "Not provided":
                    for part in reversed(parts):
                        if not re.fullmatch(r"(?:19|20)\d{2}", part):
                            school = part
                            break

            field = line
            in_split = re.split(r"\bin\b", line, maxsplit=1, flags=re.IGNORECASE)
            if len(in_split) > 1:
                field_candidate = in_split[1].strip()
                field_candidate = re.split(r"[-|,]", field_candidate)[0].strip()
                if field_candidate:
                    field = field_candidate

            education.append({
                "degree": parts[0] if parts else matched_degree.title(),
                "school": school,
                "field": field,
                "startDate": self._normalize_date(year_matches[0]) if len(year_matches) > 1 else "",
                "endDate": self._normalize_date(year_matches[-1]) if year_matches else "",
                "current": "present" in lowered or "current" in lowered,
            })

        return education

    def extract_work_experience(self, text: str) -> List[Dict[str, str]]:
        experience: List[Dict[str, str]] = []
        lines = self._get_section_lines(text, "experience") or self._clean_lines(text)

        def looks_like_title(value: str) -> bool:
            low = value.lower()
            return any(kw in low for kw in self.job_title_keywords) or bool(re.search(r"\b(lead|senior|junior|principal|head|manager|engineer|developer|analyst|consultant|specialist)\b", low))

        def looks_like_company(value: str) -> bool:
            low = value.lower()
            if any(token in low for token in ["university", "college", "school"]):
                return False
            return bool(re.search(r"\b(ltd|limited|pty|inc|llc|corp|company|solutions|technologies|technology|studio|group|consulting|systems|bank|hospital|municipal|department)\b", low)) or (len(value.split()) <= 5 and value[0].isupper())

        def parse_dates(block: str) -> Tuple[str, str, bool]:
            match = self.date_range_pattern.search(block)
            if match:
                start = self._normalize_date(match.group(1))
                end = self._normalize_date(match.group(2) or "")
                current = (match.group(2) or "").lower() in {"present", "current"}
                return start, end, current

            years = re.findall(r"(?:19|20)\d{2}", block)
            if years:
                start = self._normalize_date(years[0])
                end = self._normalize_date(years[-1]) if len(years) > 1 else ""
                current = "present" in block.lower() or "current" in block.lower()
                return start, end, current

            return "", "", False

        idx = 0
        while idx < len(lines) and len(experience) < 8:
            line = lines[idx].strip("-•\u2022\t ")
            if not line:
                idx += 1
                continue

            lookahead = " ".join(lines[idx: idx + 5])
            lower = line.lower()
            next_line = lines[idx + 1].strip("-•\u2022\t ") if idx + 1 < len(lines) else ""
            next_next = lines[idx + 2].strip("-•\u2022\t ") if idx + 2 < len(lines) else ""

            has_signal = looks_like_title(line) or looks_like_title(next_line) or bool(self.date_range_pattern.search(lookahead))
            if not has_signal:
                idx += 1
                continue

            title = line if looks_like_title(line) else (next_line if looks_like_title(next_line) else (next_next if looks_like_title(next_next) else line))

            company = ""
            for candidate in [line, next_line, next_next, lines[idx + 3].strip("-•\u2022\t ") if idx + 3 < len(lines) else ""]:
                if candidate and (" at " in candidate.lower() or looks_like_company(candidate) or " | " in candidate or "," in candidate):
                    company = re.split(r"\s+at\s+", candidate, flags=re.IGNORECASE)[-1].split("|")[0].split(",")[0].strip()
                    break

            bullets = []
            for candidate in lines[idx: idx + 5]:
                clean = candidate.strip()
                if clean and (clean.startswith("-") or clean.startswith("•") or clean.startswith("*") or len(clean) > 45):
                    bullets.append(clean.lstrip("-•* "))

            block = " ".join(lines[idx: idx + 5])
            start_date, end_date, current = parse_dates(block)

            description = " ".join(bullets[:3])[:400] if bullets else block[:300]
            experience.append({
                "title": title[:140] or "Not provided",
                "company": company[:140] or "Not provided",
                "location": None,
                "startDate": start_date,
                "endDate": end_date,
                "current": current,
                "description": description,
            })

            idx += 3

        # Deduplicate by title/company/date combination
        deduped: List[Dict[str, str]] = []
        seen = set()
        for item in experience:
            key = (item.get("title", "").lower(), item.get("company", "").lower(), item.get("startDate", ""), item.get("endDate", ""))
            if key in seen:
                continue
            seen.add(key)
            deduped.append(item)
        return deduped

    def extract_certifications(self, text: str) -> List[Dict[str, str]]:
        certifications = []
        for line in self._get_section_lines(text, "certifications")[:10]:
            if len(line) < 3:
                continue
            certifications.append({
                "name": line[:120],
                "issuer": "Not provided",
                "date": None,
            })
        return certifications

    def extract_professional_summary(self, text: str) -> Optional[str]:
        lines = self._get_section_lines(text, "summary")
        if lines:
            summary = " ".join(lines[:4]).strip()
            return summary or None

        # Fallback to the first substantial paragraph after the header/contact block.
        cleaned = self._clean_lines(text)
        for line in cleaned[:12]:
            lowered = line.lower()
            if any(token in lowered for token in ["@", "linkedin", "github", "phone", "email"]):
                continue
            if len(line) >= 40 and not re.search(r"\b(?:education|experience|skills|projects|certifications)\b", lowered):
                return line.strip()
        return None

    def extract_cover_letter(self, text: str) -> Optional[str]:
        lines = self._get_section_lines(text, "cover_letter")
        if lines:
            return "\n".join(lines[:20]).strip() or None

        # If the uploaded file bundles a letter and a CV, keep the opening paragraph as a best-effort cover letter.
        cleaned = self._clean_lines(text)
        if not cleaned:
            return None

        end_headers = {h for values in self.section_headers.values() for h in values if h not in self.section_headers.get("cover_letter", [])}
        collected: List[str] = []
        for line in cleaned[:30]:
            lowered = line.lower().strip()
            if any(h in lowered for h in end_headers):
                break
            if len(line) < 2:
                continue
            collected.append(line)
            if len(" ".join(collected)) > 1200:
                break

        return "\n".join(collected).strip() or None

    def extract_projects(self, text: str) -> List[Dict[str, str]]:
        projects: List[Dict[str, str]] = []
        lines = self._get_section_lines(text, "projects")
        if not lines:
            lines = [line for line in self._clean_lines(text) if any(token in line.lower() for token in ["project", "github", "portfolio", "app", "built", "developed", "created"])]

        current_project = None
        for raw_line in lines[:20]:
            line = raw_line.strip("-•\u2022\t ")
            if not line:
                continue

            links = self.url_pattern.findall(line)
            is_header_like = bool(re.match(r"^[A-Z][A-Za-z0-9 &/\-]{2,80}$", line)) and len(line.split()) <= 10
            if is_header_like or line.lower().startswith(("project", "portfolio")):
                if current_project:
                    projects.append(current_project)
                current_project = {
                    "name": line[:140],
                    "link": links[0] if links else "",
                    "type": "publication" if "publication" in line.lower() else "project",
                    "description": "",
                }
                continue

            if current_project is None:
                current_project = {
                    "name": line[:140],
                    "link": links[0] if links else "",
                    "type": "publication" if "publication" in line.lower() else "project",
                    "description": "",
                }
            else:
                if links and not current_project.get("link"):
                    current_project["link"] = links[0]
                desc = current_project.get("description", "")
                current_project["description"] = (desc + " " + line).strip()[:400]

        if current_project:
            projects.append(current_project)

        deduped = []
        seen = set()
        for item in projects:
            key = item.get("name", "").lower().strip()
            if not key or key in seen:
                continue
            seen.add(key)
            deduped.append(item)
        return deduped

    def extract_languages(self, text: str) -> List[Dict[str, str]]:
        results = []
        lines = self._get_section_lines(text, "languages") or self._clean_lines(text)

        for line in lines[:20]:
            low = line.lower()
            matched = [lang for lang in self.language_keywords if lang in low]
            if not matched:
                continue

            proficiency = "Not provided"
            if any(token in low for token in ["native", "mother tongue"]):
                proficiency = "Native"
            elif any(token in low for token in ["fluent", "advanced"]):
                proficiency = "Fluent"
            elif any(token in low for token in ["intermediate", "conversational"]):
                proficiency = "Intermediate"
            elif any(token in low for token in ["basic", "beginner"]):
                proficiency = "Basic"

            for lang in matched:
                results.append({"name": lang.title(), "proficiency": proficiency})

        unique = {}
        for item in results:
            key = item["name"]
            if key not in unique or unique[key] == "Not provided":
                unique[key] = item["proficiency"]
        return [{"name": name, "proficiency": prof} for name, prof in unique.items()]

    def _cross_check_profile(self, extracted: Dict[str, object], text: str) -> Dict[str, object]:
        lowered_text = (text or "").lower()
        text_norm = lowered_text.replace(" ", "")
        verified = []
        ambiguous = []
        missing = []

        checks = {
            "name": extracted.get("full_name"),
            "email": extracted.get("email"),
            "phone": extracted.get("phone"),
            "location": extracted.get("location"),
            "title": extracted.get("title"),
        }

        for field, value in checks.items():
            if not value:
                missing.append(field)
                continue

            value_norm = str(value).lower().replace(" ", "")
            if value_norm in text_norm:
                verified.append(field)
            else:
                ambiguous.append(field)

        if not extracted.get("education"):
            missing.append("education")
        if not extracted.get("work_experience"):
            missing.append("work_experience")
        if not extracted.get("skills"):
            missing.append("skills")

        return {
            "verified_fields": verified,
            "ambiguous_fields": ambiguous,
            "missing_fields": missing,
            "needs_review": bool(ambiguous or missing),
            "review_notes": (
                "Please review highlighted fields before submitting profile updates."
                if (ambiguous or missing)
                else "All core profile fields were validated against resume text."
            ),
            "processed_at": datetime.now(timezone.utc).isoformat(),
        }

    def parse(self, file_content: bytes, file_type: str) -> Dict:
        text = self.extract_text(file_content, file_type)
        if not text:
            return {"success": False, "error": "Could not extract text from file"}

        full_name = self.extract_name(text)
        first_name, last_name = self.split_name(full_name)
        skills = self.extract_skills(text)
        soft_skills = sorted({s.title() for s in self.soft_skill_keywords if s in text.lower()})

        response = {
            "success": True,
            "resume_text": text,
            "full_name": full_name,
            "first_name": first_name,
            "last_name": last_name,
            "professional_summary": self.extract_professional_summary(text),
            "cover_letter": self.extract_cover_letter(text),
            "email": self.extract_email(text),
            "phone": self.extract_phone(text),
            "location": self.extract_location(text),
            "title": self.extract_title(text),
            "skills": skills,
            "soft_skills": soft_skills,
            "experience_years": self.extract_experience_years(text),
            "education": self.extract_education(text),
            "work_experience": self.extract_work_experience(text),
            "certifications": self.extract_certifications(text),
            "projects": self.extract_projects(text),
            "languages": self.extract_languages(text),
        }
        response["validation"] = self._cross_check_profile(response, text)
        return response
